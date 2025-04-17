import os
import time
import re
from datetime import datetime
import random
import sys
import tkinter as tk
from tkinter import filedialog

try:
    from PyPDF2 import PdfReader
    from PIL import Image
    from PIL.ExifTags import TAGS
    import docx
    from mutagen.mp3 import MP3
    from mutagen.mp4 import MP4
except ImportError:
    print("\n  Required libraries missing. Install using:")
    print("  pip install PyPDF2 Pillow python-docx mutagen")
    sys.exit(1)

def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')

def get_file_path():
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    
    print("\n  Please select a file to extract metadata from...")
    
    file_types = [
        ('All Supported Files', '*.pdf;*.jpg;*.jpeg;*.png;*.gif;*.bmp;*.tiff;*.webp;*.docx;*.mp3;*.mp4;*.m4a;*.wav;*.flac;*.aac;*.mov;*.avi;*.mkv'),
        ('PDF Files', '*.pdf'),
        ('Image Files', '*.jpg;*.jpeg;*.png;*.gif;*.bmp;*.tiff;*.webp'),
        ('Document Files', '*.docx'),
        ('Audio Files', '*.mp3;*.m4a;*.wav;*.flac;*.aac'),
        ('Video Files', '*.mp4;*.mov;*.avi;*.mkv')
    ]
    
    file_path = filedialog.askopenfilename(
        title="Select File for Metadata Extraction",
        filetypes=file_types
    )
    
    if file_path:
        return file_path
    else:
        print("  No file selected. Returning to menu...")
        time.sleep(2)
        return None

def extract_pdf_metadata(file_path):
    metadata = {}
    try:
        with open(file_path, 'rb') as f:
            pdf = PdfReader(f)
            info = pdf.metadata
            
            if info:
                for key in info:
                    clean_key = key
                    if key.startswith('/'):
                        clean_key = key[1:]
                    metadata[clean_key] = str(info[key])
            
            metadata['Pages'] = len(pdf.pages)
            
            if len(pdf.pages) > 0:
                try:
                    first_page_text = pdf.pages[0].extract_text()
                    if first_page_text:
                        preview = first_page_text[:150]
                        if len(first_page_text) > 150:
                            preview += "..."
                        metadata['Text Preview'] = preview
                except:
                    pass
    except Exception as e:
        metadata['Error'] = str(e)
    
    return metadata

def extract_image_metadata(file_path):
    metadata = {}
    try:
        with Image.open(file_path) as img:
            metadata['Format'] = img.format
            metadata['Mode'] = img.mode
            metadata['Size'] = f"{img.width}x{img.height} pixels"
            
            if hasattr(img, '_getexif') and img._getexif():
                exif = img._getexif()
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8')
                        except:
                            value = str(value)
                    metadata[tag] = str(value)
            
            if hasattr(img, 'info'):
                for key, value in img.info.items():
                    if isinstance(value, (str, int, float)):
                        metadata[key] = value
    except Exception as e:
        metadata['Error'] = str(e)
    
    return metadata

def extract_docx_metadata(file_path):
    metadata = {}
    try:
        doc = docx.Document(file_path)
        
        core_props = doc.core_properties
        if core_props:
            metadata['Title'] = core_props.title or 'None'
            metadata['Author'] = core_props.author or 'None'
            metadata['Created'] = str(core_props.created) if core_props.created else 'None'
            metadata['Modified'] = str(core_props.modified) if core_props.modified else 'None'
            metadata['Last Modified By'] = core_props.last_modified_by or 'None'
            metadata['Revision'] = str(core_props.revision) or 'None'
            
        metadata['Paragraphs'] = len(doc.paragraphs)
        metadata['Sections'] = len(doc.sections)
        
        word_count = 0
        for para in doc.paragraphs:
            word_count += len(para.text.split())
        metadata['Word Count'] = word_count
        
        if doc.paragraphs and doc.paragraphs[0].text:
            preview = doc.paragraphs[0].text[:150]
            if len(doc.paragraphs[0].text) > 150:
                preview += "..."
            metadata['Text Preview'] = preview
    except Exception as e:
        metadata['Error'] = str(e)
    
    return metadata

def extract_audio_metadata(file_path):
    metadata = {}
    try:
        if file_path.lower().endswith('.mp3'):
            audio = MP3(file_path)
            metadata['Length'] = f"{int(audio.info.length // 60)}:{int(audio.info.length % 60):02d}"
            metadata['Bitrate'] = f"{int(audio.info.bitrate / 1000)} kbps"
            metadata['Sample Rate'] = f"{audio.info.sample_rate} Hz"
            
            if hasattr(audio, 'tags') and audio.tags:
                for key in audio.tags.keys():
                    metadata[key] = str(audio.tags[key])
        
        elif file_path.lower().endswith(('.m4a', '.mp4')):
            audio = MP4(file_path)
            metadata['Length'] = f"{int(audio.info.length // 60)}:{int(audio.info.length % 60):02d}"
            metadata['Bitrate'] = f"{int(audio.info.bitrate / 1000)} kbps"
            
            for key, value in audio.tags.items():
                metadata[key] = str(value)
    except Exception as e:
        metadata['Error'] = str(e)
    
    return metadata

def extract_video_metadata(file_path):
    metadata = {}
    try:
        video = MP4(file_path)
        metadata['Length'] = f"{int(video.info.length // 60)}:{int(video.info.length % 60):02d}"
        metadata['Bitrate'] = f"{int(video.info.bitrate / 1000)} kbps" if hasattr(video.info, 'bitrate') else 'Unknown'
        
        for key, value in video.tags.items():
            metadata[key] = str(value)
    except Exception as e:
        metadata['Error'] = str(e)
    
    return metadata

def get_file_extension(file_path):
    _, ext = os.path.splitext(file_path)
    return ext.lower()

def metadata_extractor():
    clear_screen()
    print("\n  === METADATA EXTRACTOR ===")
    print("  Supported formats: PDF, Images, DOCX, Audio and Video files")
    
    file_path = get_file_path()
    if not file_path:
        return
    
    print(f"\n  Extracting metadata from: {file_path}")
    print("  " + "="*50)
    
    extension = get_file_extension(file_path)
    metadata = {}
    
    if extension in ['.pdf']:
        print("  File Type: PDF")
        metadata = extract_pdf_metadata(file_path)
    
    elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
        print(f"  File Type: Image ({extension.upper()[1:]})")
        metadata = extract_image_metadata(file_path)
    
    elif extension in ['.docx']:
        print("  File Type: Microsoft Word Document")
        metadata = extract_docx_metadata(file_path)
    
    elif extension in ['.mp3', '.m4a', '.wav', '.flac', '.aac']:
        print(f"  File Type: Audio ({extension.upper()[1:]})")
        metadata = extract_audio_metadata(file_path)
    
    elif extension in ['.mp4', '.mov', '.avi', '.mkv', '.webm']:
        print(f"  File Type: Video ({extension.upper()[1:]})")
        metadata = extract_video_metadata(file_path)
    
    else:
        print(f"  File Type: {extension.upper()[1:] if extension else 'Unknown'}")
        print("  Metadata extraction not supported for this file type.")
        time.sleep(3)
        return
    
    if 'Error' in metadata:
        print(f"\n  Error extracting metadata: {metadata['Error']}")
    else:
        print("\n  Extracted Metadata:")
        for key, value in metadata.items():
            if isinstance(value, bytes) or len(str(value)) > 100:
                value = f"{str(value)[:97]}..."
            print(f"  {key}: {value}")
    
    print("\n  File System Metadata:")
    print(f"  File Size: {os.path.getsize(file_path):,} bytes")
    print(f"  Created: {datetime.fromtimestamp(os.path.getctime(file_path))}")
    print(f"  Last Modified: {datetime.fromtimestamp(os.path.getmtime(file_path))}")
    print(f"  Last Accessed: {datetime.fromtimestamp(os.path.getatime(file_path))}")
    
    input("\n  Press Enter to continue...")

if __name__ == "__main__":
    metadata_extractor() 